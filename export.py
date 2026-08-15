"""export.py — 导出模块：Word 分析报告 + CSV 原始数据

从 gen_export_demo.py 验证过的逻辑抽取，供 API 路由复用。
"""
import csv
import datetime
import functools
import io
import os
import re
import tempfile
import threading
import time

import matplotlib
matplotlib.use("Agg")  # 无界面后端，服务器环境必需
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# matplotlib pyplot 全局状态非线程安全：FastAPI 线程池并发导出可能错图/偶发异常
# （回归修复）。所有绘图函数经 _plot_locked 串行化，且 finally 关闭全部 Figure 防泄漏。
_PLOT_LOCK = threading.Lock()

# Word COM 单实例锁（回归修复：原 _refresh_fields_docx 与 _convert_to_pdf 各持一把锁，
# 并发导出时两个线程同时 Dispatch Word → RPC server is busy/文件锁冲突）
_WORD_LOCK = threading.Lock()


def _parse_share(s) -> float | None:
    """从份额字符串稳健提取数字（数据准确性：一处实现，全局复用）

    支持：'18.2%' / '46.5%（2026年Q1）' / '1,234.5%' / '18.2% (2026)' / 全角％
    范围值（'3-5%'）取中点。解析失败返回 None（不抛异常、不产生错误数字）。
    """
    if s is None:
        return None
    t = str(s).replace("％", "%").replace(",", "")
    m = re.search(r"(\d+(?:\.\d+)?)\s*[-~至]\s*(\d+(?:\.\d+)?)", t)
    if m:  # 范围取中点
        return (float(m.group(1)) + float(m.group(2))) / 2
    m = re.search(r"(\d+(?:\.\d+)?)", t)
    return float(m.group(1)) if m else None


def _plot_locked(func):
    """装饰器：串行化 matplotlib 绘图 + 异常安全关闭所有 Figure"""
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        with _PLOT_LOCK:
            try:
                return func(*args, **kwargs)
            finally:
                plt.close("all")
    return wrapper

# 中文字体：运行时探测可用 CJK 字体（Linux/CI 无微软雅黑时回退，防图表豆腐块）
def _pick_cjk_font() -> list:
    """按优先级探测本机可用中文字体；找不到时警告并回退"""
    import logging
    import matplotlib.font_manager as fm
    candidates = ["Microsoft YaHei", "SimHei", "Noto Sans CJK SC", "Noto Sans CJK",
                  "WenQuanYi Zen Hei", "PingFang SC", "Source Han Sans SC"]
    available = {f.name for f in fm.fontManager.ttflist}
    chosen = [c for c in candidates if c in available]
    if not chosen:
        logging.warning("未找到中文字体（已探测 %s），图表中文可能显示为方块", "、".join(candidates))
    return chosen


plt.rcParams["font.sans-serif"] = _pick_cjk_font() or ["DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False

# Word 报告字体（学术论文规范）：
#   正文：宋体 小四（12pt）
#   标题：黑体，一级四号（14pt）> 二级小四（12pt）> 三级五号（10.5pt）——大小递减分级
FONT_BODY = "宋体"      # 正文
FONT_HEADING = "黑体"   # 标题
FONT_SIZES = {
    "Title": (FONT_HEADING, Pt(22)),      # 封面大标题 二号（22pt）
    "Heading 1": (FONT_HEADING, Pt(14)),  # 一级标题 四号（14pt）
    "Heading 2": (FONT_HEADING, Pt(12)),  # 二级标题 小四（12pt）
    "Heading 3": (FONT_HEADING, Pt(10.5)),  # 三级标题 五号（10.5pt）
    "Normal": (FONT_BODY, Pt(12)),        # 正文 小四（12pt）
    "List Bullet": (FONT_BODY, Pt(12)),   # 列表 小四
}


def _set_font_style(style, font_name: str, size: Pt, bold: bool = False) -> None:
    """设置段落样式的字体（含 eastAsia，否则中文字符回退宋体导致大小不一）

    必须清除 w:asciiTheme/w:eastAsiaTheme 等 theme 属性——theme 字体优先于
    直接指定的 rFonts，不清除则设置不生效（Heading 默认用 majorHAnsi 主题字体）。
    """
    style.font.name = font_name
    style.font.size = size
    style.font.bold = bold
    # 关键：中文字体必须写到 rPr/w:eastAsia，python-docx 只设 name 对中文不生效
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    # 清除 theme 属性（majorHAnsi/majorEastAsia 等），否则优先于直接指定
    for attr in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def _apply_doc_fonts(doc: Document) -> None:
    """统一整份文档的字体与段间距（学术论文规范）：
    - 正文宋体小四 + 标题黑体四号
    - 章节标题段前 24pt / 段后 12pt（拉开章节间隔）
    - 小节标题段前 15.6pt / 段后 6pt（参考实训报告排版）
    - 正文段后 3pt 微间隔，避免整块挤在一起
    """
    spacing = {
        "Title": (Pt(0), Pt(6)),
        "Heading 1": (Pt(24), Pt(12)),   # 章节：前 2 行 / 后 1 行
        "Heading 2": (Pt(15.6), Pt(6)),  # 小节：前约 1.3 行（实训报告同款）/ 后 0.5 行
        "Heading 3": (Pt(12), Pt(4)),
        "Normal": (Pt(0), Pt(3)),
        "List Bullet": (Pt(0), Pt(2)),
    }
    for name, (font_name, size) in FONT_SIZES.items():
        try:
            _set_font_style(doc.styles[name], font_name, size,
                            bold=name in ("Title", "Heading 1", "Heading 2"))
            sb, sa = spacing.get(name, (Pt(0), Pt(0)))
            doc.styles[name].paragraph_format.space_before = sb
            doc.styles[name].paragraph_format.space_after = sa
        except KeyError:
            continue


def _force_runs_font(doc: Document) -> None:
    """兜底：遍历所有段落 run 强制设置字体（模板渲染/样式继承不完全时仍生效）

    样式表定义对 Word 大多数情况有效，但 docxtpl 渲染的模板段落可能带直接格式
    （direct formatting），此时 run 级强制设置能覆盖，保证全文字体一致。
    """
    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        is_heading = style_name.startswith("Heading") or style_name == "Title"
        font_name = FONT_HEADING if is_heading else FONT_BODY
        for run in p.runs:
            run.font.name = font_name
            # 标题 run 保留样式字号（不覆盖），正文 run 无字号时给默认
            if run.font.size is None and not is_heading:
                run.font.size = Pt(12)
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn("w:eastAsia"), font_name)
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
    # 表格单元格文字也统一（正文宋体）
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = FONT_BODY
                        rpr = run._element.get_or_add_rPr()
                        rfonts = rpr.get_or_add_rFonts()
                        rfonts.set(qn("w:eastAsia"), FONT_BODY)
                        rfonts.set(qn("w:ascii"), FONT_BODY)
                        rfonts.set(qn("w:hAnsi"), FONT_BODY)


def _disable_spellcheck(doc: Document) -> None:
    """隐藏拼写/语法检查（文档级设置，随文档走）

    报告含大量英文品牌/术语（AirPods/Huawei/UN Comtrade 等），Word 拼写检查
    会标红波浪线。关键设置（不是 proofState——那只是"已检查"标记，打开时
    Word 仍会重新检查标红）：
    - w:hideSpellingErrors：仅隐藏此文档的拼写错误（谁打开都不显示波浪线）
    - w:hideGrammaticalErrors：同上，语法错误
    - w:proofState clean：辅助标记
    """
    from docx.oxml import OxmlElement

    settings = doc.settings.element
    # 删除已有同名元素再重插（防重复）
    for tag in ("w:hideSpellingErrors", "w:hideGrammaticalErrors", "w:proofState"):
        for el in settings.findall(qn(tag)):
            settings.remove(el)
    # 插到 settings 开头（schema 顺序：proofState 在前，hide 类在后）
    for tag, val in (("w:hideSpellingErrors", "true"), ("w:hideGrammaticalErrors", "true")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        settings.insert(0, el)
    ps = OxmlElement("w:proofState")
    ps.set(qn("w:spelling"), "clean")
    ps.set(qn("w:grammar"), "clean")
    settings.insert(0, ps)


def _prevent_table_split(doc: Document) -> None:
    """表格行禁止跨页断开（w:cantSplit）

    学术论文规范：表格不能从中间被分页切断。给所有表格行加 cantSplit——
    行放不下时整行移到下一页，而不是被切成两半。
    """
    from docx.oxml import OxmlElement

    for tbl in doc.tables:
        for row in tbl.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                cant = OxmlElement("w:cantSplit")
                tr_pr.append(cant)


def _add_page_numbers(doc: Document) -> None:
    """页脚加页码（居中：第 X 页 / 共 Y 页，PAGE/NUMPAGES 域）"""
    from docx.oxml import OxmlElement

    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_field(text: str) -> None:
        r = p.add_run()
        r.font.size = Pt(9)
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = text
        sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
        ph = OxmlElement("w:t"); ph.text = "1"
        end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
        for el in (b, instr, sep, ph, end):
            r._element.append(el)

    run = p.add_run("第 ")
    run.font.size = Pt(9)
    _add_field(" PAGE ")
    run = p.add_run(" 页 / 共 ")
    run.font.size = Pt(9)
    _add_field(" NUMPAGES ")
    run = p.add_run(" 页")
    run.font.size = Pt(9)


def _add_toc_field(doc: Document, anchor: object = None) -> None:
    """生成学术论文式目录（成品化：点线页码 + 跳转，不依赖 Word 打开时更新）

    必须在文档内容全部生成后调用（目录需要知道每个标题的页码域）：
    - 为每个 Heading 1 标题加书签
    - 在 anchor 段落后生成目录条目：标题文字 + 右对齐点线制表位 + PAGEREF 页码域
    - 点线制表位（dot leader）是学术论文目录标准样式；PAGEREF 支持 Ctrl+点击跳转
    """
    from docx.oxml import OxmlElement

    def _add_bookmark(p, bm_id: int, name: str) -> None:
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"), str(bm_id))
        bm_start.set(qn("w:name"), name)
        bm_end = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"), str(bm_id))
        p._p.insert(0, bm_start)
        p._p.append(bm_end)

    def _add_pageref(entry, bm_name: str) -> None:
        """条目 run 后追加 PAGEREF 页码域（tab 触发点线 + 页码）"""
        r = entry.add_run("\t")._element
        b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' PAGEREF {bm_name} \\h '
        sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
        ph = OxmlElement("w:t"); ph.text = "0"
        end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
        for el in (b, instr, sep, ph, end):
            r.append(el)

    # 1. 给所有 Heading 1 标题加书签（"目录"自身跳过）
    bm_id = 1
    headings = []
    for p in doc.paragraphs:
        if p.style.name != "Heading 1" or p.text.strip() == "目录":
            continue
        name = f"_Toc_{bm_id}"
        _add_bookmark(p, bm_id, name)
        headings.append((p.text.strip(), name))
        bm_id += 1

    # 2. 生成目录条目（正序：lxml addnext 插到上一条之后）
    from docx.text.paragraph import Paragraph as _Paragraph
    # 内部锚点链接（w:anchor）不需要 relationship 声明——只有外部链接（w:r:id）才需要
    last_el = anchor._p  # 从 anchor 段落元素开始
    for title, bm_name in headings:
        new_el = last_el.makeelement(qn("w:p"), {})
        last_el.addnext(new_el)  # 插到 last_el 之后
        entry = _Paragraph(new_el, anchor._parent)
        # 回归修复：alignment=1 是 CENTER，页码会在 15.5cm 处居中；RIGHT=2 才对齐点线
        entry.paragraph_format.tab_stops.add_tab_stop(Cm(15.5), alignment=2, leader=1)
        # 目录文字用 w:hyperlink 包裹（单击直接跳转，无需 Ctrl+点击）
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("w:anchor"), bm_name)  # 锚点 = 标题书签名
        r_el = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "微软雅黑")
        rfonts.set(qn("w:hAnsi"), "微软雅黑")
        rfonts.set(qn("w:eastAsia"), "微软雅黑")
        rpr.append(rfonts)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "123C5C")  # 深海蓝，区分普通文字
        rpr.append(color)
        r_el.append(rpr)
        t = OxmlElement("w:t")
        t.text = title
        t.set(qn("xml:space"), "preserve")
        r_el.append(t)
        hl.append(r_el)
        entry._p.append(hl)
        # PAGEREF 页码域（点线 + 页码）
        _add_pageref(entry, bm_name)
        last_el = new_el

    # settings.xml 加 updateFields（必须按 w:settings 的 schema 顺序插入，append 到末尾可能被 Word 忽略）
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        # w:settings 子元素顺序：w:zoom, w:embedSystemFonts, w:defaultTabStop, w:updateFields, ...
        # 插到第一个非 zoom 的元素前，保证在 w:defaultTabStop 之后
        anchor_el = None
        for child in settings:
            tag = child.tag.split('}')[-1]
            if tag not in ("zoom", "embedSystemFonts", "characterSpacingControl", "defaultTabStop"):
                anchor_el = child
                break
        if anchor_el is not None:
            anchor_el.addprevious(upd)
        else:
            settings.append(upd)


@_plot_locked
def build_trend_chart(trend: dict) -> io.BytesIO:
    """生成趋势折线图 PNG（内存流），供 Word 报告嵌入"""
    # 按年份排序（键可能是 str/int）；兼容 {year: {value, weight}} 与 {year: float} 两种结构
    years = sorted(trend.keys(), key=lambda k: int(k))
    def _val(y):
        v = trend[y]
        return v["value"] if isinstance(v, dict) else v
    values = [_val(y) / 1e8 for y in years]  # 亿美元

    # 宽幅布局，给标注留足空间（tight_layout 防裁切）
    fig, ax = plt.subplots(figsize=(10, 5))
    fig.subplots_adjust(left=0.1, right=0.95, top=0.88, bottom=0.15)
    ax.plot(years, values, marker="o", linewidth=2.2, color="#2e5bff")
    ax.fill_between(years, values, alpha=0.12, color="#2e5bff")
    ax.set_title("出口贸易金额趋势（亿美元）", fontsize=13)
    ax.set_xlabel("年份", fontsize=11)
    ax.set_ylabel("亿美元", fontsize=11)
    ax.grid(True, alpha=0.3)
    # 顶部留 15% 余量，防止峰值标注文字超出绘图区
    vmin, vmax = ax.get_ylim()
    ax.set_ylim(vmin, vmax * 1.15)
    for x, y in zip(years, values):
        # 标注放数据点下方（避免顶部溢出被裁切；统一向下永不出界）
        ax.annotate(f"{y:.2f}", (x, y), textcoords="offset points",
                    xytext=(0, -14), fontsize=10, ha="center")
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def build_executive_summary(product: str, target: str, year: str, stats: dict,
                            analysis: dict, total_value: float) -> str:
    """生成执行摘要（报告开头）：关键数字 + AI 一句话总结 + 数据来源标注"""
    lines = []
    if stats:
        # 数据区间（Citation：数字可溯源）；单年显示"X年"，多年显示"X-Y"
        fy, ly = stats.get("first_year"), stats.get("last_year")
        y_range = f"{fy}-{ly}" if fy and ly and fy != ly else (f"{fy}年" if fy else year)
        lines.append(f"• 总出口额: {total_value / 1e8:.2f} 亿美元（{y_range}，UN Comtrade）")
        if stats.get("cagr_pct") is not None:
            lines.append(f"• 年复合增长率: {stats['cagr_pct']}%（{y_range}）")
        if stats.get("peak_year"):
            lines.append(f"• 峰值年份: {stats['peak_year']}（{y_range} 区间内）")
        if stats.get("change_over_period_pct") is not None:
            lines.append(f"• 期末较期初变化: {stats['change_over_period_pct']:.1f}%（{y_range}）")
    if analysis.get("overview"):
        lines.append(f"• AI 总结: {analysis['overview']}")
    lines.append(f"• 数据来源: UN Comtrade 公共 API，报告生成于 {datetime.date.today().isoformat()}")
    return "\n".join(lines) if lines else f"（{product} → {target} {year}，暂无摘要数据）"


def _is_valid_pdf(path: str) -> bool:
    """PDF 有效性校验：文件存在、非 0 字节、%PDF 魔数开头（防半成品文件）"""
    try:
        if not os.path.exists(path) or os.path.getsize(path) == 0:
            return False
        with open(path, "rb") as f:
            return f.read(4) == b"%PDF"
    except OSError:
        return False


def finalize_docx(buf: io.BytesIO, as_pdf: bool = False) -> tuple:
    """报告收尾（共用）：写临时文件 → COM 更新域/修表格跨页 → 可选转 PDF → 返回

    返回 (buf, fmt)：fmt 为实际生成格式（'docx'/'pdf'），调用方按 fmt 定 media_type 和文件名。
    - docx: COM 更新 PAGEREF 页码、页脚 PAGE、表格防切分，补写拼写检查隐藏
    - pdf: 在上一步基础上 Word/LibreOffice 导出 PDF；转换失败/结果无效降级返回 docx
    - COM 全部失败：原样返回输入 buf（fmt 按请求，docx 域靠用户打开时更新）
    """
    import logging
    import os
    import tempfile
    # mkstemp 原子创建唯一临时文件（时间戳在同毫秒并发时可能撞名覆盖）
    fd, tmp_path = tempfile.mkstemp(suffix=".docx", prefix="_tp_export_")
    os.close(fd)
    try:
        with open(tmp_path, "wb") as f:
            f.write(buf.getvalue())
        _refresh_fields_docx(tmp_path)
        read_path = tmp_path
        fmt = "docx"
        if as_pdf:
            _convert_to_pdf(tmp_path)
            pdf_path = tmp_path.replace(".docx", ".pdf")
            if _is_valid_pdf(pdf_path):
                read_path = pdf_path
                fmt = "pdf"
            else:
                logging.warning("PDF 转换结果无效或缺失，降级返回 docx")
        with open(read_path, "rb") as f:
            return io.BytesIO(f.read()), fmt
    except Exception:
        # COM 全部失败：返回原始 docx（域靠用户打开时自动更新）；PDF 请求强制降级 docx，
        # 避免"docx 内容 + .pdf 后缀 + application/pdf"的损坏文件
        logging.exception("报告收尾处理失败，返回原始 docx")
        return buf, "docx"
    finally:
        for p in (tmp_path, tmp_path.replace(".docx", ".pdf")):
            try:
                os.remove(p)
            except OSError:
                pass


def _convert_to_pdf(docx_path: str) -> None:
    """docx → pdf（同目录同名 .pdf）。Word COM 优先，LibreOffice headless 回退
    （Linux/Docker/无 Word 环境），两者都失败记日志（回归修复：此前完全静默且
    Linux 下 PDF 永远降级 docx——LibreOffice 白装）。
    """
    import logging
    import shutil
    import subprocess

    with _WORD_LOCK:
        # 路径 1：Word COM（Windows + 已装 Word）
        try:
            import win32com.client
            # DispatchEx 强制新建独立实例（回归修复：Dispatch 会连接用户已打开的
            # Word，随后 Quit() 会关掉用户文档，有数据丢失风险）
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                word.DisplayAlerts = 0
                doc = word.Documents.Open(docx_path)
                doc.SaveAs2(docx_path.replace(".docx", ".pdf"), FileFormat=17)
                doc.Close()
            finally:
                word.Quit()
            return
        except Exception as e:
            logging.warning("Word COM 转 PDF 失败（尝试 LibreOffice 回退）: %s", e)
        # 路径 2：LibreOffice headless（跨平台后备）
        try:
            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            if soffice is None:
                logging.warning("未找到 soffice/libreoffice，PDF 转换不可用（将降级 docx）")
                return
            out_dir = os.path.dirname(os.path.abspath(docx_path))
            subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
                timeout=120, capture_output=True,
            )
        except Exception as e:
            logging.warning("LibreOffice 转 PDF 失败（将降级 docx）: %s", e)


def build_word_report(product: str, target: str, year: str, hs_code: str,
                      rows: list, ai: dict, hs_description: str = "",
                      stats: dict | None = None, analysis: dict | None = None,
                      landscape: dict | None = None,
                      market_ctx: dict | None = None,
                      matrix: list | None = None,
                      background: dict | None = None,
                      competitiveness: dict | None = None,
                      reporter: str = "中国") -> io.BytesIO:
    """生成贸易数据 Word 报告（与市场分析同套规范：封面/目录/字体/页码/表格防切）

    章节：封面 → 目录 → 一、执行摘要 → 二、出口趋势（图）→ 三、数据总览（表）
    → 四、出口大国对比（矩阵）→ 五、竞争格局 → 六、目标市场消费环境
    → 七、原始数据（表）→ 八、AI 市场分析 → 附录：数据来源。
    """
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor

    total_value = sum(r.get("primaryValue") or 0 for r in rows)
    total_wgt = sum(r.get("netWgt") or 0 for r in rows)
    hs_desc = f"（{hs_description}）" if hs_description else ""

    # 趋势图 PNG（≥2 年才生成）；用 summarize_trend 逐年累加，与执行摘要 stats 同口径
    chart_buf = None
    from trade import summarize_trend, get_latest_year
    trend_map = summarize_trend(rows)
    if len(trend_map) >= 2:
        chart_buf = build_trend_chart(trend_map)

    doc = Document()
    _apply_doc_fonts(doc)
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(12)
    NAVY = RGBColor(0x12, 0x3C, 0x5C)
    ACCENT = RGBColor(0xC4, 0x45, 0x2C)

    def _hr(space_before: bool = True):
        p = doc.add_paragraph()
        if space_before:
            p.paragraph_format.space_before = Pt(6)
        ppr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "123C5C")
        pbdr.append(bottom)
        ppr.append(pbdr)
        return p

    def _h(text, level=1, blank_before=False):
        if level == 1 and blank_before:
            for _ in range(2):
                blank = doc.add_paragraph()
                # 空行段落加极小 run（2pt 字）：被推到页首时只是微距而非空两行
                br = blank.add_run(" ")
                br.font.size = Pt(2)
                blank.paragraph_format.space_before = Pt(0)
                blank.paragraph_format.space_after = Pt(0)
                blank.paragraph_format.line_spacing = Pt(2)
                blank.paragraph_format.keepNext = True
        return doc.add_heading(text, level=level)

    def _p(text="", bold=False, indent=True):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)
        r = p.add_run(text)
        r.bold = bold
        return p

    # ===== 封面 =====
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_before = Pt(12)
    br = brand.add_run("TRADEPILOT AI  ·  EXPORT INTELLIGENCE")
    br.font.size = Pt(11)
    br.font.color.rgb = NAVY
    br.bold = True
    _hr()
    t = doc.add_heading(f"{product}出口贸易分析报告", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(24)
    for run in t.runs:
        run.font.color.rgb = NAVY
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_before = Pt(6)
    sr = st.add_run(f"目标市场：{target}  ·  出口国：中国  ·  HS{hs_code}{hs_desc}")
    sr.font.size = Pt(14)
    sr.font.color.rgb = ACCENT
    sr.bold = True
    doc.add_paragraph()
    info_lines = [
        f"数据来源：UN Comtrade 联合国商品贸易数据库",
        f"生成日期：{datetime.date.today().isoformat()}  ·  报告编号：TP-{datetime.date.today().strftime('%Y%m%d')}-{target}-HS{hs_code}",
        "统计指标由程序精确计算 · AI 仅作解读 · 数据可溯源",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _hr(space_before=True)
    # 核心数据速览（封面 KPI 区标题：非章节标题，用 Normal 加粗避免混进目录）
    kpi_title = doc.add_paragraph()
    kpi_title.paragraph_format.space_before = Pt(12)
    ktr = kpi_title.add_run("核心数据速览")
    ktr.font.size = Pt(14)
    ktr.bold = True
    kpi_rows = []
    # 回归修复：多年报告封面 KPI 用年份区间标注（原写死单年 year，与多年总量对不上）；
    # 与执行摘要同口径：stats.first_year-last_year
    if stats and stats.get("first_year") and stats.get("last_year") and stats["first_year"] != stats["last_year"]:
        year_label_kpi = f"{stats['first_year']}-{stats['last_year']}"
    else:
        year_label_kpi = str(year)
    kpi_rows.append((f"对{target}出口总额（{year_label_kpi}）", f"{total_value / 1e8:.2f} 亿美元"))
    if total_wgt:
        kpi_rows.append(("出口净重", f"{total_wgt / 1e6:.2f} 千吨"))
    if stats:
        if stats.get("cagr_pct") is not None:
            kpi_rows.append(("年复合增长率 CAGR", f"{stats['cagr_pct']}%"))
        if stats.get("peak_year"):
            kpi_rows.append(("峰值年份", f"{stats['peak_year']}"))
    kpi_rows.append(("数据记录数", f"{len(rows)} 条"))
    if kpi_rows:
        kpi_tbl = doc.add_table(rows=len(kpi_rows), cols=2)
        kpi_tbl.style = "Light Grid Accent 1"
        for i, (a, b) in enumerate(kpi_rows):
            kpi_tbl.rows[i].cells[0].text = a
            kpi_tbl.rows[i].cells[1].text = b
            for cell in kpi_tbl.rows[i].cells:
                for cp in cell.paragraphs:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _hr()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("TradePilot AI · Export Intelligence")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ===== 目录（学术论文式：点线页码 + 单击跳转）=====
    doc.add_page_break()
    toc_title = _h("目录", 1)
    doc.add_page_break()

    # ===== 一、执行摘要 =====
    _h("一、执行摘要", 1, blank_before=True)
    summary_text = build_executive_summary(
        product, target, year,
        stats or {}, analysis or {},
        total_value,
    )
    for line in summary_text.split("\n"):
        _p(line)

    # ===== 二、出口趋势（图）=====
    _h("二、出口趋势", 1, blank_before=True)
    if chart_buf:
        _p(f"{reporter}对{target}出口 {product}（HS {hs_code}）出口额变化趋势，单位：亿美元：")
        doc.add_picture(chart_buf, width=Cm(14))
        _p(f"数据来源：UN Comtrade 公共 API（HS {hs_code}）", indent=False)
    else:
        _p("（单年数据，无趋势图）")

    # ===== 三、数据总览表 =====
    _h("三、数据总览", 1, blank_before=True)
    tbl = doc.add_table(rows=3, cols=3)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "指标", "数值", "单位"
    tbl.rows[1].cells[0].text = "贸易金额"
    tbl.rows[1].cells[1].text = f"{total_value:,.0f}"
    tbl.rows[1].cells[2].text = "美元"
    tbl.rows[2].cells[0].text = "净重"
    tbl.rows[2].cells[1].text = f"{total_wgt:,.0f}"
    tbl.rows[2].cells[2].text = "公斤"
    if stats:
        _h("统计指标（程序精确计算）", 2)
        if stats.get("change_over_period_pct") is not None:
            _p(f"• 区间 {stats['first_year']}-{stats['last_year']} 期末较期初变化 {stats['change_over_period_pct']:.1f}%", indent=False)
        if stats.get("max_swing_year") is not None:
            _p(f"• 最大单年波动：{stats['max_swing_year']} 年 {stats['max_swing_pct']}%", indent=False)
        prices = stats.get("unit_prices") or []
        if prices:
            _p("• 单价趋势：" + "; ".join(f"{p['year']}年 {p['price']:.2f} 美元/公斤" for p in prices), indent=False)
            # 单价柱状图（量价结构一眼看清）
            _p("单价（美元/公斤）逐年变化：")
            _add_bar_chart(doc, {p["year"]: p["price"] for p in prices}, "出口单价趋势", "美元/公斤")

    # ===== 四、出口大国对比（饼图主角 + 表格作证）=====
    _h("四、出口大国对比", 1, blank_before=True)
    if matrix:
        _p(f"该品类对{target}的出口大国竞争格局：")
        # 饼图先放：份额结构一眼看清
        share_labels = [m.get("country", "") for m in matrix if m.get("market_share") is not None]
        share_vals = [m["market_share"] for m in matrix if m.get("market_share") is not None]
        if share_labels:
            _p(f"各出口国占 {target} 市场进口份额（%，<4% 合并为其他）：")
            _add_pie_chart(doc, share_labels, share_vals, f"{product} 出口国份额结构", raw_values=share_vals)
        # 饼图"其他"拆解 + 全量明细表（作证饼图 + 延伸分析）
        _h("份额明细与「其他」拆解", 2)
        _p("饼图中的「其他」包含以下出口国——份额虽小但增速与单价各不相同，值得单独观察：", indent=False)
        mtbl = doc.add_table(rows=1 + len(matrix), cols=6)
        mtbl.style = "Light Grid Accent 1"
        for j, head in enumerate(["出口国", "最新出口(亿美元)", "占市场进口份额", "5年CAGR", "单价($/kg)", "判断"]):
            mtbl.rows[0].cells[j].text = head
        for i, m in enumerate(matrix, 1):
            mtbl.rows[i].cells[0].text = str(m.get("country", ""))
            mtbl.rows[i].cells[1].text = f"{m.get('export_value', 0) / 1e8:.2f}" if m.get("export_value") else "—"
            mtbl.rows[i].cells[2].text = f"{m['market_share']}%" if m.get("market_share") is not None else "—"
            mtbl.rows[i].cells[3].text = f"{m['cagr_pct']:+.1f}%" if m.get("cagr_pct") is not None else "—"
            mtbl.rows[i].cells[4].text = f"${m['unit_price']:.2f}" if m.get("unit_price") is not None else "—"
            mtbl.rows[i].cells[5].text = str(m.get("verdict", ""))
        # 解读：谁在涨谁在跌（基于矩阵数据）
        _h("竞争态势解读", 2)
        rising = [m for m in matrix if (m.get("cagr_pct") or 0) > 5]
        falling = [m for m in matrix if (m.get("cagr_pct") or 0) < -2]
        leader = matrix[0] if matrix else {}
        if leader.get("country"):
            # market_share 可为 None（目标市场无进口数据），防 "None%" 渲染（回归修复）
            share_txt = f"{leader['market_share']}%" if leader.get("market_share") is not None else "暂无数据"
            _p(f"• {leader['country']}以 {leader.get('export_value', 0) / 1e8:.2f} 亿美元居首（占市场进口 {share_txt}），"
               f"CAGR {(leader.get('cagr_pct') or 0):+.1f}%（{leader.get('verdict', '')}）。", indent=False)
        if rising:
            names = "、".join(m["country"] for m in rising[:3])
            cagrs = "、".join("{:+.1f}%".format(m["cagr_pct"] or 0) for m in rising[:3])
            _p(f"• 上升方：{names}（CAGR {cagrs}）——"
               f"这些出口国份额在扩大，是{leader.get('country', '中国')}的主要追赶者。", indent=False)
        if falling:
            names = "、".join(m["country"] for m in falling[:3])
            cagrs = "、".join("{:+.1f}%".format(m["cagr_pct"] or 0) for m in falling[:3])
            _p(f"• 下滑方：{names}（CAGR {cagrs}）——"
               f"份额收缩，竞争压力相对缓解。", indent=False)
    else:
        _p("（出口大国对比数据不足）")

    # ===== 五、竞争格局（饼图主角 + 表格作证）=====
    _h("五、竞争格局", 1, blank_before=True)
    if landscape and landscape.get("top_brands"):
        brands = landscape["top_brands"]
        _p(f"{landscape.get('product_category', product)} 龙头品牌竞争格局（来源：{landscape.get('_source', '行业检索')}）。")
        # 品牌份额饼图先放（解析 share 里的数字，用模块级 _parse_share 统一处理）
        pie_labels = []
        pie_vals = []
        for b in brands[:8]:
            v = _parse_share(b.get("share"))
            if v is not None and v > 0:
                pie_labels.append(b.get("name", ""))
                pie_vals.append(v)
        if pie_labels:
            _p(f"龙头品牌市场份额结构（<4% 合并为其他）：")
            _add_pie_chart(doc, pie_labels, pie_vals, f"{product} 品牌份额", raw_values=pie_vals)
        # 份额排名表（作证饼图 + 地位说明）
        _h("品牌份额排名", 2)
        _p("饼图份额对应的品牌全量明细——排名、份额与市场地位：", indent=False)
        btbl = doc.add_table(rows=1 + len(brands), cols=3)
        btbl.style = "Light Grid Accent 1"
        for j, head in enumerate(["品牌", "市场份额", "市场地位"]):
            btbl.rows[0].cells[j].text = head
        for i, b in enumerate(brands, 1):
            btbl.rows[i].cells[0].text = str(b.get("name", ""))
            btbl.rows[i].cells[1].text = str(b.get("share", ""))
            btbl.rows[i].cells[2].text = str(b.get("position", ""))
        if landscape.get("shift_reasons"):
            _h("格局变动原因", 2)
            for r in landscape["shift_reasons"]:
                _p(f"• {r}", indent=False)
        if landscape.get("chain_insight"):
            _h("产业链洞察", 2)
            _p(f"• {landscape['chain_insight']}", indent=False)
    else:
        _p("（竞争格局数据不足）")

    # ===== 六、目标市场消费环境 =====
    _h("六、目标市场消费环境", 1, blank_before=True)
    if market_ctx and market_ctx.get("available"):
        env = []
        if market_ctx.get("gdp"):
            env.append(f"GDP {market_ctx['gdp'] / 1e12:.2f} 万亿美元")
        if market_ctx.get("population"):
            env.append(f"人口 {market_ctx['population'] / 1e8:.2f} 亿")
        if market_ctx.get("gdp_per_capita"):
            env.append(f"人均 GDP {market_ctx['gdp_per_capita']:,.0f} 美元")
        _p(f"{target} 经济环境（World Bank）：{'、'.join(env)}")
        if market_ctx.get("gdp_per_capita"):
            pc = market_ctx["gdp_per_capita"]
            level = "高收入市场（消费力强，支撑中高端产品溢价）" if pc > 30000 else (
                "中等收入市场（性价比敏感）" if pc > 10000 else "发展中市场（价格驱动）")
            _p(f"• 需求判断：人均 GDP {pc:,.0f} 美元 → {level}。", indent=False)
        if market_ctx.get("population"):
            _p(f"• 人口 {market_ctx['population'] / 1e8:.2f} 亿：人口规模决定市场容量上限。", indent=False)
        if background and background.get("global_trade_growth"):
            _p(f"• 宏观背景：全球贸易增长预测 {background['global_trade_growth']}（{background.get('_source', 'WTO')}）。", indent=False)
    else:
        _p("（目标市场经济数据不足）")

    # ===== 七、驱动因素分析（CPI/科技出口 + 需求/供给/竞争三侧）=====
    _h("七、驱动因素分析", 1, blank_before=True)
    _p("驱动出口与销量变化的因素可分为需求侧、供给侧、竞争侧：")
    # 需求侧：CPI 通胀趋势（World Bank 免费 API）→ 趋势图
    if market_ctx and market_ctx.get("available") and market_ctx.get("iso3"):
        iso3 = market_ctx["iso3"]
        cpi_series = {}
        try:
            from market_data import get_worldbank_series
            years5 = list(range(get_latest_year() - 4, get_latest_year() + 1))
            cpi_series = get_worldbank_series(iso3, "cpi", years5)
        except Exception:
            pass
        if cpi_series:
            _h("需求侧：通胀与消费环境", 2)
            _p(f"{target} 通胀率（CPI 年变化 %，World Bank 官方数据）近 5 年趋势：")
            _add_line_chart(doc, cpi_series, "CPI 通胀率变化", "%")
            sorted_cpi = sorted(cpi_series.items())
            latest_cpi = sorted_cpi[-1][1] if sorted_cpi else None
            if latest_cpi is not None:
                level = "低通胀（消费环境稳定，利于可选消费支出）" if latest_cpi < 3 else (
                    "温和通胀（消费略有压力）" if latest_cpi < 5 else "高通胀（消费承压，可选消费萎缩）")
                _p(f"• 最新通胀 {latest_cpi:.1f}%：{level}。", indent=False)
            if len(sorted_cpi) >= 2:
                first_cpi = sorted_cpi[0][1]
                last_cpi = sorted_cpi[-1][1]
                if first_cpi and first_cpi > 3 and last_cpi < first_cpi:
                    _p(f"• 通胀从 {first_cpi:.1f}% 回落至 {last_cpi:.1f}%：购买力修复，"
                       f"对消费电子产品需求是利好信号。", indent=False)
    # 供给侧：高科技出口占比（出口能力结构）
    if market_ctx and market_ctx.get("high_tech_exports") is not None:
        _h("供给侧：出口能力结构", 2)
        _p(f"• {target} 高科技出口占制成品出口 {market_ctx['high_tech_exports']:.1f}%"
           f"（World Bank）——该市场自身科技产业基础，决定对进口消费电子的依赖度。", indent=False)
    if market_ctx and market_ctx.get("mobile") is not None:
        _p(f"• 每百人手机订阅 {market_ctx['mobile']:.0f} 部：移动设备渗透率支撑智能硬件需求。", indent=False)
    # 驱动因素数据表（因素 / 数据 / 影响）
    factor_rows = [("驱动因素", "数据", "影响方向")]
    if market_ctx and market_ctx.get("gdp_per_capita"):
        factor_rows.append(("人均 GDP（消费力）", f"{market_ctx['gdp_per_capita']:,.0f} 美元", "高收入市场支撑中高端溢价"))
    if market_ctx and market_ctx.get("cpi") is not None:
        factor_rows.append(("通胀率 CPI", f"{market_ctx['cpi']:.1f}%", "低通胀利于可选消费支出"))
    if market_ctx and market_ctx.get("high_tech_exports") is not None:
        factor_rows.append(("高科技出口占比", f"{market_ctx['high_tech_exports']:.1f}%", "科技产业基础 → 进口依赖度"))
    if competitiveness and competitiveness.get("tc") is not None:
        factor_rows.append(("贸易竞争力 TC", f"{competitiveness['tc']}",
                            "强则出口主导，弱则进口依赖"))
    if competitiveness and competitiveness.get("market_share") is not None:
        factor_rows.append(("占市场进口份额", f"{competitiveness['market_share']}%", "现有渗透率 = 增长基数"))
    if stats and stats.get("cagr_pct") is not None:
        factor_rows.append(("出口 CAGR", f"{stats['cagr_pct']}%", "出口动能方向"))
    if landscape and landscape.get("top_brands"):
        factor_rows.append(("龙头品牌份额", f"{landscape['top_brands'][0].get('share', '')}", "市场集中度决定进入难度"))
    if len(factor_rows) > 1:
        ftbl = doc.add_table(rows=len(factor_rows), cols=3)
        ftbl.style = "Light Grid Accent 1"
        for i, (a, b, c) in enumerate(factor_rows):
            ftbl.rows[i].cells[0].text = a
            ftbl.rows[i].cells[1].text = b
            ftbl.rows[i].cells[2].text = c
            if i == 0:
                for cell in ftbl.rows[0].cells:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True
        _p()
        _p("上表数据来源：UN Comtrade / World Bank / 行业检索。")

    # ===== 八、原始数据表 =====
    _h("八、原始数据（UN Comtrade）", 1, blank_before=True)
    raw_tbl = doc.add_table(rows=1 + len(rows), cols=5)
    raw_tbl.style = "Light Grid Accent 1"
    for j, head in enumerate(["年份", "流向", "HS编码", "金额(美元)", "净重(公斤)"]):
        raw_tbl.rows[0].cells[j].text = head
    for i, r in enumerate(rows, 1):
        raw_tbl.rows[i].cells[0].text = str(r.get("refYear"))
        raw_tbl.rows[i].cells[1].text = "出口"
        raw_tbl.rows[i].cells[2].text = str(r.get("cmdCode"))
        raw_tbl.rows[i].cells[3].text = f"{float(r.get('primaryValue') or 0):,.0f}"
        raw_tbl.rows[i].cells[4].text = f"{float(r.get('netWgt') or 0):,.0f}"

    # ===== 九、AI 市场分析 =====
    _h("九、AI 市场分析", 1, blank_before=True)
    ms = ai.get("market_size") or {}
    gt = ai.get("growth_trend") or {}
    risks = ai.get("risks") or []
    up = ai.get("user_profile") or {}
    _h("市场规模", 2)
    # 防重复：AI 的 value 常自带"（2026年估算）"，再追加会双份（渲染 bug 修复）
    ms_value = str(ms.get("value", "未知"))
    ms_year = str(ms.get("year", ""))
    ms_suffix = f"（{ms_year}年估算）" if ms_year and "估算" not in ms_value else ""
    _p(f"{ms_value}{ms_suffix}")
    _h("增长趋势", 2)
    _p(f"CAGR {gt.get('cagr', '未知')}，{gt.get('forecast_years', '')}")
    if gt.get("description"):
        _p(gt["description"])
    _h("用户画像", 2)
    _p(f"年龄区间: {up.get('age_range', '')} | 收入水平: {up.get('income_level', '')}")
    _h("风险分析", 2)
    for r in risks:
        if isinstance(r, dict):
            _p(f"• {r.get('type')}（{r.get('level')}）: {r.get('description')}", indent=False)
    _h("AI 总结", 2)
    _p(ai.get("summary", ""))

    # ===== 附录：数据来源 =====
    _h("附录：数据来源与说明", 1, blank_before=True)
    src_rows = [("数据维度", "来源", "说明")]
    src_rows.append(("出口贸易数据", "UN Comtrade 联合国商品贸易统计数据库", "HS 编码口径，公共 API 实时查询"))
    src_rows.append(("统计指标", "程序精确计算", "CAGR / 区间变化 / 最大波动 / 单价趋势"))
    stbl = doc.add_table(rows=len(src_rows), cols=3)
    stbl.style = "Light Grid Accent 1"
    for i, (a, b, c) in enumerate(src_rows):
        stbl.rows[i].cells[0].text = a
        stbl.rows[i].cells[1].text = b
        stbl.rows[i].cells[2].text = c
        if i == 0:
            for cell in stbl.rows[0].cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
    _p()
    _p("免责声明：本报告数据来源与统计口径见附录，市场估算部分仅供参考，实际决策请以官方统计为准。"
       "贸易竞争力指数（TC）与市场份额基于 UN Comtrade 出口（FOB 口径）与进口（CIF 口径）数据计算，"
       "两口径存在约 5-10% 的系统性差异，属贸易统计惯例。")

    # 收尾统一：页码 / 拼写检查 / 表格防切 / 目录 / 字体
    _add_page_numbers(doc)
    _disable_spellcheck(doc)
    _prevent_table_split(doc)
    _add_toc_field(doc, toc_title)
    _force_runs_font(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _refresh_fields_docx(path: str) -> None:
    """用 Word COM 打开 docx 并更新所有域（PAGEREF 页码/页脚 PAGE），保存后返回

    顺带修复表格跨页：检测每个表格首行/末行页码，跨页时表格前插入分页符
    （整表移到下一页，不切断）。仅 Windows + 已装 Word 时可用；失败静默。
    用线程锁串行化：Word COM 单实例，并发调用会冲突。
    """
    with _WORD_LOCK:
        try:
            import win32com.client
            # DispatchEx 强制新建独立实例（回归修复：Dispatch 会连接用户已打开的
            # Word，随后 Quit() 会关掉用户文档，有数据丢失风险）
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            try:
                word.DisplayAlerts = 0  # 防兼容性/恢复弹窗卡死
                doc = word.Documents.Open(path)
                doc.Repaginate()
                # 清理页首空行段落（章节空行被推到新页首时删除，避免页首空两行）
                for _ in range(3):
                    doc.Repaginate()
                    cleaned = False
                    for pg in range(1, doc.ComputeStatistics(2) + 1):  # wdStatisticPages
                        try:
                            first_para = doc.Range(
                                doc.GoTo(1, 1, 1, pg).Start,  # wdGoToPage
                                doc.GoTo(1, 1, 1, pg).Start
                            ).Paragraphs(1)
                        except Exception:
                            continue
                        txt = first_para.Range.Text.strip()
                        # 空行特征：无文字（可能含空格/极小字 run）且不是表格
                        if txt in ("", " ") and first_para.Range.Tables.Count == 0:
                            if first_para.Range.Information(3) == pg:
                                first_para.Range.Delete()
                                cleaned = True
                                break
                    if not cleaned:
                        break
                # 检测并修复表格跨页（最多 3 轮，插入分页符后分页变化需重新检查）
                for _ in range(3):
                    doc.Repaginate()
                    fixed = False
                    for tbl in doc.Tables:
                        try:
                            first_pg = tbl.Rows(1).Range.Information(3)  # wdActiveEndPageNumber
                            last_pg = tbl.Rows(tbl.Rows.Count).Range.Information(3)
                        except Exception:
                            continue
                        if first_pg != last_pg:
                            # 表格跨页：在表格第一行前插入分页符（整表下移一页）
                            tbl.Rows(1).Range.InsertBreak(7)  # wdPageBreak
                            fixed = True
                            break  # 分页已变，重新检查
                    if not fixed:
                        break
                doc.Repaginate()
                doc.Fields.Update()
                for section in doc.Sections:
                    for f in section.Footers:
                        if f.Range.Fields.Count:
                            f.Range.Fields.Update()
                # 关闭拼写/语法检查（COM 重写 settings.xml 会清掉，保存后补写）
                try:
                    doc.SpellingChecked = True
                    doc.GrammarChecked = True
                except Exception:
                    pass
                doc.Save()
                doc.Close()
            finally:
                word.Quit()
            # COM 保存后补写文档级"隐藏拼写错误"设置（COM 重写 settings.xml 会清掉）
            from docx import Document as _Doc
            _doc = _Doc(path)
            _disable_spellcheck(_doc)
            _doc.save(path)
        except Exception:
            pass  # 无 Word/COM 失败：交给 updateFields 打开时更新


def build_market_report(product: str, country: str, ai: dict,
                        trade_evidence: dict | None = None,
                        competitiveness: dict | None = None,
                        background: dict | None = None,
                        landscape: dict | None = None,
                        market_context: dict | None = None,
                        trend_series: dict | None = None) -> io.BytesIO:
    """生成市场分析 Word 报告（20 页版：原始数据摘要 + 经济/贸易/竞争力图表 + 竞争格局 + 驱动因素）"""
    from market_data import get_worldbank_series, COUNTRY_ISO3
    from trade import get_latest_year
    from docx.shared import RGBColor
    import logging

    doc = Document()
    # 统一字体：正文宋体小四 + 标题黑体（学术论文规范）
    _apply_doc_fonts(doc)
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(12)

    iso3 = COUNTRY_ISO3.get(country, "")
    latest = get_latest_year()
    years5 = list(range(latest - 4, latest + 1))

    def _h(text, level=1, blank_before=False):
        """章节标题：level 1 且 blank_before 时前插 2 个空行段落，与上一章节明显隔开

        封面/目录页开头的标题不插（blank_before=False，页首无需空行）。
        空行段落加 keepNext，保证空行+标题不拆分跨页（避免页底留半页空白）。
        """
        if level == 1 and blank_before:
            for _ in range(2):
                blank = doc.add_paragraph()
                # 空行段落加极小 run（2pt 字）：被推到页首时只是微距而非空两行
                br = blank.add_run(" ")
                br.font.size = Pt(2)
                blank.paragraph_format.space_before = Pt(0)
                blank.paragraph_format.space_after = Pt(0)
                blank.paragraph_format.line_spacing = Pt(2)
                blank.paragraph_format.keepNext = True  # 与下一段（标题）同页
        return doc.add_heading(text, level=level)

    def _p(text="", bold=False, indent=True):
        """正文段落：中文首行缩进 2 字符（默认），列表项/标题行不缩进"""
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)  # 2 字符（12pt 字号 × 2）
        r = p.add_run(text)
        r.bold = bold
        return p

    # ===== 封面（学术论文式：品牌行 + 装饰线 + 大标题 + 信息块 + 核心速览）=====
    from docx.oxml import OxmlElement
    from docx.shared import RGBColor
    NAVY = RGBColor(0x12, 0x3C, 0x5C)
    ACCENT = RGBColor(0xC4, 0x45, 0x2C)

    def _hr(space_before: bool = True):
        """装饰线：段落底边框（skill 建议：横线用段落边框，不用表格）"""
        p = doc.add_paragraph()
        if space_before:
            p.paragraph_format.space_before = Pt(6)
        ppr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")      # 1.5pt 线宽
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "123C5C")  # 深海蓝
        pbdr.append(bottom)
        ppr.append(pbdr)
        return p

    # 顶部品牌行
    brand = doc.add_paragraph()
    brand.alignment = WD_ALIGN_PARAGRAPH.CENTER
    brand.paragraph_format.space_before = Pt(12)
    br = brand.add_run("TRADEPILOT AI  ·  EXPORT INTELLIGENCE")
    br.font.size = Pt(11)
    br.font.color.rgb = NAVY
    br.bold = True
    _hr()

    # 大标题（居中）
    t = doc.add_heading(f"{product}市场分析报告", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_before = Pt(24)
    for run in t.runs:
        run.font.color.rgb = NAVY
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.paragraph_format.space_before = Pt(6)
    sr = st.add_run(f"目标市场：{country}  ·  出口国：中国")
    sr.font.size = Pt(14)
    sr.font.color.rgb = ACCENT
    sr.bold = True

    # 信息块（居中段落：数据来源 / 生成日期 / 报告编号）
    doc.add_paragraph()
    info_lines = [
        f"数据来源：UN Comtrade · World Bank · Tavily 行业检索",
        f"生成日期：{datetime.date.today().isoformat()}  ·  报告编号：TP-{datetime.date.today().strftime('%Y%m%d')}-{country}",
        "统计指标由程序精确计算 · AI 仅作解读 · 数据可溯源",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line)
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _hr(space_before=True)

    # 核心数据速览（封面 KPI 区标题：非章节标题，用 Normal 加粗避免混进目录）
    kpi_title = doc.add_paragraph()
    kpi_title.paragraph_format.space_before = Pt(12)
    ktr = kpi_title.add_run("核心数据速览")
    ktr.font.size = Pt(14)
    ktr.bold = True
    kpi_rows = []
    if trade_evidence and trade_evidence.get("trend"):
        trend = trade_evidence["trend"]
        years = sorted(trend.keys())
        last_y, last_v = years[-1], trend[str(years[-1])]
        kpi_rows.append((f"对{country}出口额（{last_y} 年）", f"{last_v} 亿美元"))
    if competitiveness and competitiveness.get("tc") is not None:
        kpi_rows.append(("贸易竞争力指数 TC", f"{competitiveness['tc']}"))
        if competitiveness.get("market_share") is not None:
            kpi_rows.append(("占该国市场进口份额", f"{competitiveness['market_share']}%"))
    if market_context and market_context.get("gdp_per_capita"):
        kpi_rows.append(("人均 GDP", f"{market_context['gdp_per_capita']:,.0f} 美元"))
    if market_context and market_context.get("gdp"):
        kpi_rows.append(("GDP", f"{market_context['gdp'] / 1e12:.2f} 万亿美元"))
    if kpi_rows:
        kpi_tbl = doc.add_table(rows=len(kpi_rows), cols=2)
        kpi_tbl.style = "Light Grid Accent 1"
        for i, (a, b) in enumerate(kpi_rows):
            kpi_tbl.rows[i].cells[0].text = a
            kpi_tbl.rows[i].cells[1].text = b
            for cell in kpi_tbl.rows[i].cells:
                for cp in cell.paragraphs:
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 底部品牌
    _hr()
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("TradePilot AI · Export Intelligence")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 目录（学术论文式：点线页码 + 跳转，独立成页）
    doc.add_page_break()
    toc_title = _h("目录", 1)
    doc.add_page_break()  # 目录后分页（正文从新页开始）
    # 目录条目在文档全部生成后填充（末尾统一生成，见 _build_toc_entries）

    # ===== 一、执行摘要（关键数字速览表）=====
    _h("一、执行摘要", 1, blank_before=True)
    _p("本报告数据来源：联合国商品贸易数据库（UN Comtrade）提供出口贸易数据，"
       "世界银行（World Bank）提供经济环境数据，行业检索提供竞争格局与宏观背景。"
       "所有统计指标（CAGR、贸易竞争力指数、市场出口份额）由程序精确计算，AI 仅作解读，数据可溯源。")
    es = ai.get("executive_summary") or {}
    if es.get("background"):
        _p(f"背景：{es['background']}")
    if es.get("key_findings"):
        _h("核心发现", 2)
        for f in es["key_findings"]:
            _p(f"• {f}", indent=False)
    if es.get("challenges"):
        _h("主要挑战", 2)
        for c in es["challenges"]:
            _p(f"• {c}", indent=False)
    if es.get("recommendation"):
        _p(f"总体建议：{es['recommendation']}")
    if es.get("data_points"):
        _h("关键数据点", 2)
        for d in es["data_points"]:
            _p(f"• {d}", indent=False)

    # ===== 二、市场环境（经济趋势图）=====
    _h("二、市场环境", 1, blank_before=True)
    if iso3:
        gdp_series = get_worldbank_series(iso3, "gdp", years5)
        pc_series = get_worldbank_series(iso3, "gdp_per_capita", years5)
        if gdp_series:
            _p(f"{country} 近 5 年 GDP 变化（World Bank 官方数据，单位：万亿美元）：")
            _add_bar_chart(doc, gdp_series, "GDP 变化趋势", "万亿美元", divisor=1e12)
        if pc_series:
            _p(f"{country} 近 5 年人均 GDP 变化（World Bank 官方数据，单位：美元）：")
            _add_line_chart(doc, pc_series, "人均 GDP 变化趋势", "美元")
        _h("经济环境解读", 2)
        env = []
        if market_context and market_context.get("gdp"):
            env.append(f"GDP {market_context['gdp'] / 1e12:.2f} 万亿美元")
        if market_context and market_context.get("population"):
            env.append(f"人口 {market_context['population'] / 1e8:.2f} 亿")
        if market_context and market_context.get("gdp_per_capita"):
            env.append(f"人均 GDP {market_context['gdp_per_capita']:,.0f} 美元")
        if env:
            _p(f"{country} 最新经济数据：{'，'.join(env)}（World Bank）")
        if background and background.get("summary"):
            _p(f"全球宏观背景：{background['summary']}（{background.get('_source', 'WTO')}）")
    else:
        _p("（该国家不在 World Bank 数据覆盖范围内，市场环境图表跳过）")

    # ===== 三、出口贸易（趋势图 + 汇总表）=====
    _h("三、出口贸易", 1, blank_before=True)
    if trade_evidence and trade_evidence.get("trend"):
        trend = trade_evidence["trend"]
        years = sorted(trend.keys())
        _p(f"中国对{country}出口 {product}（HS {trade_evidence.get('hs_code', '')}）逐年数据，单位：亿美元：")
        _add_line_chart(doc, {int(y): float(v) for y, v in trend.items()}, "出口额趋势", "亿美元")
        # 年度明细表（含净重与单价）
        wt = trade_evidence.get("weight_trend") or {}
        tbl = doc.add_table(rows=1 + len(years), cols=4)
        tbl.style = "Light Grid Accent 1"
        for j, head in enumerate(["年份", "出口额（亿美元）", "净重（千吨）", "同比变化"]):
            tbl.rows[0].cells[j].text = head
        prev = None
        for i, y in enumerate(years, 1):
            v = trend[str(y)]
            tbl.rows[i].cells[0].text = str(y)
            tbl.rows[i].cells[1].text = f"{v}"
            w = wt.get(str(y))
            tbl.rows[i].cells[2].text = f"{w}" if w else "—"
            if prev is not None and prev > 0:
                chg = (v - prev) / prev * 100
                tbl.rows[i].cells[3].text = f"{chg:+.1f}%"
            else:
                tbl.rows[i].cells[3].text = "—"
            prev = v
        # 数据构成解读：单价趋势 + 波动上下文（程序计算，AI 不参与算术）
        _h("数据构成解读", 2)
        unit_prices = []
        for y in years:
            v = trend[str(y)]
            w = wt.get(str(y))
            if v and w and w > 0:
                unit_prices.append((int(y), v * 1e8 / (w * 1e6)))  # 美元/公斤
        if len(unit_prices) >= 2:
            _p("单价（出口额 ÷ 净重）逐年变化，反映产品价值构成：")
            for y, up in unit_prices:
                _p(f"• {y} 年：{up:.2f} 美元/公斤", indent=False)
            first_up = unit_prices[0][1]
            last_up = unit_prices[-1][1]
            if first_up > 0:
                chg = (last_up - first_up) / first_up * 100
                trend_desc = "上升（向高价值产品结构演进）" if chg > 5 else (
                    "下降（低价值产品占比扩大或价格竞争加剧）" if chg < -5 else "平稳")
                _p(f"• 区间单价变化 {chg:+.1f}%：{trend_desc}——"
                   f"单价上升通常意味着产品向中高端升级，单价下降则可能面临价格竞争或低端产品放量。")
        # 波动上下文（峰值/谷值/最大波动）
        if len(years) >= 3:
            vals = [(int(y), trend[str(y)]) for y in years]
            peak_y, peak_v = max(vals, key=lambda x: x[1])
            trough_y, trough_v = min(vals, key=lambda x: x[1])
            swings = []
            for k in range(1, len(vals)):
                a, b = vals[k - 1][1], vals[k][1]
                if a > 0:
                    swings.append((vals[k][0], (b - a) / a * 100))
            max_swing = max(swings, key=lambda x: abs(x[1])) if swings else None
            _p(f"• 区间峰值：{peak_y} 年 {peak_v} 亿美元；谷值：{trough_y} 年 {trough_v} 亿美元。", indent=False)
            if max_swing:
                _p(f"• 最大单年波动：{max_swing[0]} 年 {max_swing[1]:+.1f}%"
                   f"（{'下滑' if max_swing[1] < 0 else '增长'}——"
                   f"需结合当年宏观事件与行业周期解读，如需求波动、供应链调整或贸易政策变化）。")
        _p()
        _p(f"数据来源：UN Comtrade 公共 API（HS {trade_evidence.get('hs_code', '')}），单价与波动为程序计算")
    else:
        _p("（该产品对目标市场暂无贸易数据）")

    # ===== 四、竞争力分析（TC + 份额）=====
    _h("四、竞争力分析", 1, blank_before=True)
    if competitiveness and competitiveness.get("tc") is not None:
        tc = competitiveness["tc"]
        share = competitiveness.get("market_share")
        exp_v = competitiveness.get("export_value", 0) / 1e8
        imp_v = competitiveness.get("import_value", 0) / 1e8
        _p(f"贸易竞争力指数（TC）= {tc}，取值 -1 到 1，越接近 1 表示出口竞争力越强。")
        _p(f"中国对{country}出口 {product}：出口 {exp_v:.2f} 亿美元 vs 进口 {imp_v:.2f} 亿美元。")
        _add_gauge_chart(doc, tc, "TC 贸易竞争力指数")
        _h("进出口结构解读", 2)
        net = exp_v - imp_v
        if exp_v > 0:
            _p(f"• 净出口 {net:+.2f} 亿美元：中国在该品类对{country}呈"
               f"{'贸易顺差' if net > 0 else '贸易逆差'}。")
        if exp_v + imp_v > 0:
            _p(f"• 进出口比 {exp_v / (exp_v + imp_v) * 100:.0f}% 出口占比："
               f"{'以出口为主导' if exp_v > imp_v else '进口依赖明显'}——"
               f"该数据反映中国在该市场的角色是{'供应方' if exp_v > imp_v else '采购方'}。")
        if share is not None:
            _p(f"• 市场出口份额 {share}%：中国产品占{country}该品类进口的比重，"
               f"即该国每进口 100 美元该品类，约 {share:.0f} 美元来自中国。")
            _p(f"• 份额解读：{'渗透率较高，进入成熟竞争期' if share >= 15 else (
                '渗透率中等，仍有扩张空间' if share >= 5 else '渗透率较低，市场拓展空间大')}。")
        # 四态数据质量标记（数据准确性红线）：rejected 时明确告知不可用，
        # suspicious 时披露口径差异；AI 不参与判定，全部来自程序 DataGate
        q = competitiveness.get("quality")
        if q == "rejected":
            _p(f"⚠️ 数据无法用于本次分析：{competitiveness.get('quality_note') or '原始数据未通过完整性校验'}", indent=False)
        elif q == "suspicious" and competitiveness.get("quality_note"):
            _p(f"⚠️ 数据质量提示：{competitiveness['quality_note']}", indent=False)
    else:
        _p("（该品类竞争力数据不足）")

    # ===== 五、竞争格局（饼图主角 + 表格作证 + 龙头财报）=====
    _h("五、竞争格局", 1, blank_before=True)
    if landscape and landscape.get("top_brands"):
        brands = landscape["top_brands"]
        _p(f"{landscape.get('product_category', product)} 龙头品牌竞争格局"
           f"（来源：{landscape.get('_source', 'Tavily 行业检索')}）。")
        # 品牌份额饼图先放（解析 share 里的数字，用模块级 _parse_share 统一处理）
        pie_labels = []
        pie_vals = []
        for b in brands[:8]:
            v = _parse_share(b.get("share"))
            if v is not None and v > 0:
                pie_labels.append(b.get("name", ""))
                pie_vals.append(v)
        if pie_labels:
            _p(f"龙头品牌市场份额结构（<4% 合并为其他）：")
            _add_pie_chart(doc, pie_labels, pie_vals, f"{product} 品牌份额", raw_values=pie_vals)
        # 份额排名表（作证饼图）
        _h("品牌份额排名", 2)
        _p("饼图份额对应的品牌全量明细——排名、份额与市场地位：", indent=False)
        tbl = doc.add_table(rows=1 + len(brands), cols=3)
        tbl.style = "Light Grid Accent 1"
        for j, head in enumerate(["品牌", "市场份额", "地位"]):
            tbl.rows[0].cells[j].text = head
        for i, b in enumerate(brands, 1):
            tbl.rows[i].cells[0].text = str(b.get("name", ""))
            tbl.rows[i].cells[1].text = str(b.get("share", ""))
            tbl.rows[i].cells[2].text = str(b.get("position", ""))
        if landscape.get("shift_reasons"):
            _h("格局变动原因", 2)
            for r in landscape["shift_reasons"]:
                _p(f"• {r}", indent=False)
        if landscape.get("chain_insight"):
            _h("产业链洞察", 2)
            _p(landscape["chain_insight"])
        # 龙头财报（结合财务画像：SEC 美股 / A 股 / 非上市公开报道）
        _h("龙头品牌财务画像", 2)
        _p("结合龙头品牌的公开财报，判断其投入能力与市场策略（数据源：SEC 财报 / 东方财富 / 公开报道）：")
        financials_available = False
        # 进程内缓存：同公司 24h 内不重查（SEC 查询 ~15 秒/次，报告每次生成都调会拖慢）
        # 阶段 4：24h TTL；回归修复：模块级锁 double-check（防并发同品牌重复慢查询）
        _fin_cache = getattr(build_market_report, "_fin_cache", None)
        if _fin_cache is None:
            _fin_cache = {}
            build_market_report._fin_cache = _fin_cache
        _fin_ts = getattr(build_market_report, "_fin_ts", None)
        if _fin_ts is None:
            _fin_ts = {}
            build_market_report._fin_ts = _fin_ts
        _fin_lock = getattr(build_market_report, "_fin_lock", None)
        if _fin_lock is None:
            _fin_lock = threading.Lock()
            build_market_report._fin_lock = _fin_lock
        try:
            from financials import get_company_financials
            # 品牌 → 财务查询名映射（兼容中英文品牌名，覆盖常见消费电子龙头）
            brand_cn = {"Apple": "苹果", "Huawei": "华为", "Xiaomi": "小米", "Sony": "索尼",
                        "Samsung": "三星", "Edifier": "漫步者", "JBL": "哈曼",
                        "Anker": "安克创新", "DJI": "大疆", "GoPro": "GoPro",
                        "Bose": "Bose", "Razer": "雷蛇", "Xiaomi": "小米"}
            for b in brands[:3]:
                fname = brand_cn.get(b.get("name", ""), b.get("name", ""))
                # 进程内缓存命中（24h 内不重查；超时重新拉取）。double-check：
                # 锁内再查一次，并发同品牌只发起一次 SEC 慢查询
                _now = time.time()
                if fname in _fin_cache and _now - _fin_ts.get(fname, 0) < 86400:
                    fin = _fin_cache[fname]
                else:
                    with _fin_lock:
                        if fname in _fin_cache and _now - _fin_ts.get(fname, 0) < 86400:
                            fin = _fin_cache[fname]
                        else:
                            try:
                                fin = get_company_financials(fname)
                                _fin_cache[fname] = fin
                                _fin_ts[fname] = time.time()
                            except Exception:
                                # 瞬时故障不写缓存（否则被掩盖 24h），本次跳过该品牌
                                logging.warning("财务画像拉取失败（本次跳过）: %s", fname)
                                continue
                if not fin or not fin.get("available"):
                    continue
                metrics = fin.get("metrics") or {}
                rev_series = metrics.get("revenue") or []
                if not rev_series:
                    continue
                financials_available = True
                latest = rev_series[-1]
                rev_val = latest.get("value")
                # 单位标注（数据准确性）：优先用 financials 模块返回的 unit 字段，
                # 回退到 source 前缀判断（SEC 美元 / A股·非上市 人民币元）
                unit = fin.get("unit", "")
                if unit == "USD":
                    rev_unit = "美元"
                elif unit == "CNY":
                    rev_unit = "人民币元"
                else:
                    rev_unit = "美元" if fin.get("source", "").startswith("SEC") else "人民币元"
                _p(f"• {fname}（份额 {b.get('share', '')}）：", indent=False)
                if rev_val:
                    _p(f"  最新营收 {rev_val / 1e8:.2f} 亿{rev_unit}（{latest.get('year', '')}年）"
                       f"——体量支撑其市场投入与研发（来源：{fin.get('source', '')}）。", indent=False)
                # 近 3 年营收趋势（扩张/收缩判断）
                if len(rev_series) >= 3:
                    first_v = rev_series[0].get("value")
                    last_v = rev_series[-1].get("value")
                    if first_v and last_v and first_v > 0:
                        chg = (last_v - first_v) / first_v * 100
                        direction = "营收扩张，有持续投入能力" if chg > 10 else (
                            "营收稳定，投入平稳" if chg > -5 else "营收收缩，投入承压")
                        _p(f"  近 {len(rev_series)} 年营收变化 {chg:+.1f}%：{direction}。", indent=False)
        except Exception:
            pass
        if not financials_available:
            _p("（所选龙头品牌的公开财报数据暂缺，跳过财务画像）", indent=False)
        if landscape.get("key_insight"):
            _h("核心洞察", 2)
            _p(landscape["key_insight"])
        _h("对出口商的启示", 2)
        top_share = None
        for b in brands:
            s = _parse_share(b.get("share"))
            if s is not None:
                top_share = max(top_share or 0, s)
        if top_share is not None:
            if top_share >= 30:
                _p(f"• 龙头品牌（最大者）份额约 {top_share:.0f}%：市场集中度较高，新进入者宜避开正面竞争，"
                   f"从细分场景（如通勤降噪、运动佩戴、价格带空档）切入。")
            elif top_share >= 15:
                _p(f"• 龙头品牌（最大者）份额约 {top_share:.0f}%：市场存在主导者，存在差异化空间，"
                   f"可在功能或价格带建立差异化定位。")
            else:
                _p(f"• 龙头品牌（最大者）份额约 {top_share:.0f}%：市场相对分散，竞争格局未固化，"
                   f"是新进入者布局的窗口期。")
        if landscape.get("shift_reasons"):
            _p(f"• 格局正在变动（{'；'.join(landscape['shift_reasons'][:2])}），"
               f"变动期往往伴随新品牌崛起的机会，建议关注变动的技术/渠道驱动因素。")
        if landscape.get("chain_insight"):
            # 完整引用产业链洞察（不截断拼接模板套话，避免"存储芯片（…）：掌握…"式怪句）
            _p(f"• 产业链洞察：{landscape['chain_insight']}", indent=False)
    else:
        _p("（竞争格局数据不足，此章节跳过）")

    # ===== 六、市场规模与增长 =====
    _h("六、市场规模与增长", 1, blank_before=True)
    ms = ai.get("market_size") or {}
    # 防重复：AI 的 value 常自带"（2026年估算）"，再追加会双份（渲染 bug 修复）
    ms_value = str(ms.get("value", "未知"))
    ms_year = str(ms.get("year", ""))
    ms_suffix = f"（{ms_year}年估算）" if ms_year and "估算" not in ms_value else ""
    _p(f"规模：{ms_value}{ms_suffix}")
    if ms.get("note"):
        _p(f"说明：{ms['note']}")
    gt = ai.get("growth_trend") or {}
    if gt.get("cagr") or gt.get("description"):
        _h("增长趋势", 2)
        if gt.get("cagr"):
            _p(f"年复合增长率（CAGR）：{gt['cagr']}（{gt.get('forecast_years', '')}）")
        if gt.get("description"):
            _p(gt["description"])
        if gt.get("key_drivers"):
            _h("关键驱动因素", 3)
            for d in gt["key_drivers"]:
                _p(f"• {d}", indent=False)

    # ===== 七、驱动因素分析（整合数据）=====
    _h("七、驱动因素分析", 1, blank_before=True)
    _p("驱动出口与销量变化的因素可分为需求侧、供给侧、竞争侧：")
    # 驱动因素数据表（因素 / 数据 / 影响方向）
    factor_rows = [("驱动因素", "数据", "影响")]
    if market_context and market_context.get("gdp_per_capita"):
        factor_rows.append(("人均 GDP（消费力）", f"{market_context['gdp_per_capita']:,.0f} 美元",
                            "高收入市场支撑中高端产品溢价"))
    if market_context and market_context.get("population"):
        factor_rows.append(("人口规模", f"{market_context['population'] / 1e8:.2f} 亿", "决定市场容量上限"))
    if market_context and market_context.get("gdp"):
        factor_rows.append(("经济总量 GDP", f"{market_context['gdp'] / 1e12:.2f} 万亿美元", "整体需求底盘"))
    if background and background.get("global_trade_growth"):
        factor_rows.append(("全球贸易增长预测", background["global_trade_growth"], "宏观景气度影响出口节奏"))
    if competitiveness and competitiveness.get("tc") is not None:
        factor_rows.append(("贸易竞争力 TC", f"{competitiveness['tc']}",
                            f"{'竞争力强，可扩张' if competitiveness['tc'] > 0.5 else '竞争力中等，需提升'}"))
    if competitiveness and competitiveness.get("market_share") is not None:
        factor_rows.append(("占市场进口份额", f"{competitiveness['market_share']}%", "现有渗透率 = 增长基数"))
    if trade_evidence and trade_evidence.get("trend"):
        trend = trade_evidence["trend"]
        years = sorted(trend.keys())
        if len(years) >= 2:
            first, last = trend[str(years[0])], trend[str(years[-1])]
            # 防 0 值（某年出口额为 0 时 pow(last/0) 崩溃）
            if first > 0 and last > 0:
                n_years = int(years[-1]) - int(years[0])
                cagr = (pow(last / first, 1 / n_years) - 1) * 100 if n_years > 0 else None
            else:
                cagr = None
            # 标注"历史"：与第六章 AI 预测 CAGR 区分（避免同报告两个 CAGR 打架）
            factor_rows.append((f"出口 CAGR（{years[0]}-{years[-1]} 历史）",
                                f"{cagr:+.1f}%" if cagr is not None else "—",
                                "出口动能方向（历史）"))
    if landscape and landscape.get("top_brands"):
        factor_rows.append(("龙头品牌份额", f"{landscape['top_brands'][0].get('share', '')}",
                            "市场集中度决定进入难度"))
    if len(factor_rows) > 1:
        ftbl = doc.add_table(rows=len(factor_rows), cols=3)
        ftbl.style = "Light Grid Accent 1"
        for i, (a, b, c) in enumerate(factor_rows):
            ftbl.rows[i].cells[0].text = a
            ftbl.rows[i].cells[1].text = b
            ftbl.rows[i].cells[2].text = c
            if i == 0:
                for cell in ftbl.rows[0].cells:
                    if cell.paragraphs[0].runs:
                        cell.paragraphs[0].runs[0].bold = True
        _p()
        _p("上表数据来源：UN Comtrade / World Bank / 行业检索。")

    _h("需求侧（经济与消费力）", 2)
    if market_context and market_context.get("gdp_per_capita"):
        _p(f"• 人均 GDP {market_context['gdp_per_capita']:,.0f} 美元：高收入市场对中高端产品需求强，"
           f"支撑 {product} 的溢价空间。")
    if market_context and market_context.get("population"):
        _p(f"• 人口 {market_context['population'] / 1e8:.2f} 亿：人口规模决定市场容量上限。", indent=False)
    if background and background.get("global_trade_growth"):
        _p(f"• 全球贸易增长预测 {background['global_trade_growth']}：宏观景气度影响整体出口节奏。", indent=False)
    _h("供给侧（出口能力）", 2)
    if competitiveness and competitiveness.get("tc") is not None:
        _p(f"• TC 指数 {competitiveness['tc']}：中国在该品类对{country}的贸易竞争力"
           f"（{'强' if competitiveness['tc'] > 0.5 else '中等'}）。")
    if trade_evidence and trade_evidence.get("trend"):
        trend = trade_evidence["trend"]
        years = sorted(trend.keys())
        if len(years) >= 2:
            first, last = trend[str(years[0])], trend[str(years[-1])]
            # 防 0 值（某年出口额为 0 时 pow(last/0) 崩溃）
            if first > 0 and last > 0:
                n_years = int(years[-1]) - int(years[0])
                cagr = (pow(last / first, 1 / n_years) - 1) * 100 if n_years > 0 else None
                _p(f"• {years[0]}-{years[-1]} 出口 CAGR {cagr:+.1f}%（历史）：反映该品类在目标市场的整体出口动能。", indent=False)
    _h("竞争侧（格局与份额）", 2)
    if landscape and landscape.get("top_brands"):
        brands = landscape["top_brands"]
        top = brands[0]
        _p(f"• 龙头 {top.get('name', '')} 份额 {top.get('share', '')}：市场集中度决定新进入者难度。", indent=False)
    if competitiveness and competitiveness.get("market_share") is not None:
        _p(f"• 中国占进口份额 {competitiveness['market_share']}%：现有渗透率是增长的基数。", indent=False)

    # ===== 八、风险分析 =====
    _h("八、风险分析", 1, blank_before=True)
    risks = ai.get("risks") or []
    if risks:
        tbl = doc.add_table(rows=1 + len(risks), cols=4)
        tbl.style = "Light Grid Accent 1"
        # 列宽分配（法规列加宽，避免"94/62/EC）、CE 认证"式截断）
        for row in tbl.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(1.5)
            row.cells[2].width = Cm(6.5)
            row.cells[3].width = Cm(5.0)
        for j, head in enumerate(["风险类型", "等级", "说明", "相关法规"]):
            tbl.rows[0].cells[j].text = head
        for i, r in enumerate(risks, 1):
            if isinstance(r, dict):
                tbl.rows[i].cells[0].text = str(r.get("type", ""))
                tbl.rows[i].cells[1].text = str(r.get("level", ""))
                tbl.rows[i].cells[2].text = str(r.get("description", ""))
                tbl.rows[i].cells[3].text = str(r.get("regulation", ""))
    else:
        _p("（风险数据不足）")

    # ===== 九、用户画像与目标客群 =====
    _h("九、用户画像与目标客群", 1, blank_before=True)
    up = ai.get("user_profile") or {}
    if up.get("age_range") or up.get("income_level"):
        _p(f"年龄区间：{up.get('age_range', '未知')}  |  收入水平：{up.get('income_level', '未知')}")
    if up.get("key_needs"):
        _h("核心需求", 2)
        for n in up["key_needs"]:
            _p(f"• {n}", indent=False)
    if up.get("buying_habits"):
        _h("购买习惯", 2)
        for b in up["buying_habits"]:
            _p(f"• {b}", indent=False)
    _h("客群细分与触达建议", 2)
    if market_context and market_context.get("gdp_per_capita"):
        _p(f"• 高端客群：人均 GDP {market_context['gdp_per_capita']:,.0f} 美元支撑，主打品质与品牌溢价，"
           f"触达渠道以品牌官网 / 高端连锁为主。")
    _p(f"• 主流客群：注重性价比与口碑，触达以电商平台（如 Amazon 等）+ 社媒内容种草为主。", indent=False)
    _p(f"• 年轻客群：偏好新潮设计与社交属性，可借助 TikTok / Instagram 短视频与 KOL 合作扩大曝光。", indent=False)
    _h("购买决策旅程", 2)
    _p(f"1. 认知阶段：通过平台搜索、社媒种草、评测视频了解 {product} 品类的头部品牌与新品。")
    _p(f"2. 比较阶段：对比价格、参数、评论口碑，重点关注降噪、续航、佩戴舒适等核心卖点。")
    _p(f"3. 决策阶段：参考电商评分与销量榜单，高收入客群倾向品牌溢价，主流客群倾向性价比。")
    _p(f"4. 复购阶段：满意的佩戴体验与售后保障驱动复购，品牌生态（如智能互联）增强粘性。")
    _h("对营销策略的启示", 2)
    _p(f"• 定价分层：参考人均消费力与竞品价位，设置引流款（入门）+ 利润款（中高端）组合。", indent=False)
    _p(f"• 内容策略：围绕核心需求制作场景化内容（通勤降噪 / 运动佩戴 / 办公会议），匹配搜索意图。", indent=False)
    _p(f"• 渠道组合：电商平台为主力，社媒种草为杠杆，本地售后为信任背书。", indent=False)

    # ===== 十、行动路线 =====
    _h("十、行动路线", 1, blank_before=True)
    ap = ai.get("action_plan") or []
    if ap:
        for i, step in enumerate(ap, 1):
            _p(f"{i}. {step}")
    if ai.get("outlook"):
        _h("市场展望", 2)
        _p(ai["outlook"])
    if ai.get("summary"):
        _h("AI 总结", 2)
        _p(ai["summary"])

    # ===== 附录：数据来源与免责声明 =====
    # 附录自然接在正文后（不强制分页，避免产生空白页）
    _h("附录：数据来源与说明", 1, blank_before=True)
    src_rows = [("数据维度", "来源", "说明")]
    src_rows.append(("出口贸易数据", "UN Comtrade 联合国商品贸易统计数据库", "HS 编码口径，公共 API 实时查询"))
    src_rows.append(("经济环境数据", "World Bank 世界银行开放数据", "GDP / 人口 / 人均 GDP / 互联网普及率"))
    src_rows.append(("宏观背景", "WTO 全球贸易展望（Tavily 检索）", "30 天增量刷新"))
    src_rows.append(("竞争格局", "Tavily 行业检索", "龙头品牌 / 市场份额 / 格局变动原因，30 天缓存"))
    src_rows.append(("统计指标", "程序精确计算", "CAGR / TC 指数 / 市场出口份额 / 同比变化"))
    stbl = doc.add_table(rows=len(src_rows), cols=3)
    stbl.style = "Light Grid Accent 1"
    for i, (a, b, c) in enumerate(src_rows):
        stbl.rows[i].cells[0].text = a
        stbl.rows[i].cells[1].text = b
        stbl.rows[i].cells[2].text = c
        if i == 0:
            for cell in stbl.rows[0].cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
    _p()
    _p("免责声明：本报告数据来源与统计口径见附录，市场估算部分（如市场规模数值）仅供参考，"
       "实际决策请以官方统计与一手调研为准。所有可溯源的统计指标均由程序计算，AI 不参与算术。")

    # 页脚页码（第 X 页 / 共 Y 页）
    _add_page_numbers(doc)
    # 关闭拼写检查（品牌/术语不标红波浪线）
    _disable_spellcheck(doc)

    # 目录条目（文档全部生成后填充：书签 + 点线页码条目）
    _add_toc_field(doc, toc_title)

    # 表格行禁止跨页断开（学术论文规范）
    _prevent_table_split(doc)

    # 兜底：run 级强制统一字体（模板直接格式/样式继承不完全时仍生效）
    _force_runs_font(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@_plot_locked
def _add_bar_chart(doc: Document, series: dict, title: str, unit: str, divisor: float = 1.0):
    """柱状图：多年数据 → matplotlib PNG → 嵌入 Word"""
    import matplotlib.pyplot as plt
    years = sorted(int(y) for y in series.keys())
    values = [series[y] / divisor for y in years]
    fig, ax = plt.subplots(figsize=(8, 4))
    fig.subplots_adjust(left=0.12, right=0.95, top=0.85, bottom=0.15)
    ax.bar([str(y) for y in years], values, color="#2e5bff", alpha=0.85)
    ax.set_title(f"{title}（{unit}）", fontsize=12)
    ax.grid(True, alpha=0.3, axis="y")
    for x, v in zip(range(len(years)), values):
        ax.annotate(f"{v:.2f}", (x, v), textcoords="offset points",
                    xytext=(0, 4), fontsize=9, ha="center")
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130)
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=Cm(14))


@_plot_locked
def _add_line_chart(doc: Document, series: dict, title: str, unit: str, divisor: float = 1.0):
    """折线图：多年数据 → matplotlib PNG → 嵌入 Word"""
    import matplotlib.pyplot as plt
    years = sorted(int(y) for y in series.keys())
    values = [series[y] / divisor for y in years]
    fig, ax = plt.subplots(figsize=(8, 4))
    fig.subplots_adjust(left=0.12, right=0.95, top=0.85, bottom=0.15)
    ax.plot([str(y) for y in years], values, marker="o", linewidth=2.2, color="#2e5bff")
    ax.fill_between(range(len(years)), values, alpha=0.12, color="#2e5bff")
    ax.set_title(f"{title}（{unit}）", fontsize=12)
    ax.grid(True, alpha=0.3)
    for x, v in zip(range(len(years)), values):
        ax.annotate(f"{v:.2f}", (x, v), textcoords="offset points",
                    xytext=(0, -14), fontsize=9, ha="center")
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130)
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=Cm(14))


@_plot_locked
def _add_gauge_chart(doc: Document, tc: float, title: str):
    """竞争力仪表图：TC 指数 → 半圆仪表 PNG"""
    import numpy as np
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(8, 3.5))
    # 半圆背景（-1 到 1）
    theta = np.linspace(0, np.pi, 100)
    x = np.cos(theta); y = np.sin(theta)
    ax.plot(x, y, color="#d0d0d0", linewidth=12, solid_capstyle="round")
    # 指示弧（0 到 TC）
    theta2 = np.linspace(np.pi, np.pi - (tc + 1) / 2 * np.pi, 50)
    x2 = np.cos(theta2); y2 = np.sin(theta2)
    ax.plot(x2, y2, color="#2e5bff" if tc >= 0 else "#c4452c", linewidth=12, solid_capstyle="round")
    ax.text(0, -0.2, f"TC = {tc:.2f}", ha="center", fontsize=16, fontweight="bold", color="#2e5bff")
    ax.text(-1.15, -0.1, "-1", ha="center", fontsize=10, color="#888")
    ax.text(1.15, -0.1, "1", ha="center", fontsize=10, color="#888")
    ax.text(0, 1.15, "竞争力弱 ← → 竞争力强", ha="center", fontsize=10, color="#666")
    ax.set_xlim(-1.3, 1.3); ax.set_ylim(-0.4, 1.3)
    ax.set_aspect("equal"); ax.axis("off")
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130)
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=Cm(13))


@_plot_locked
def _add_pie_chart(doc: Document, labels: list, values: list, title: str, raw_values: list = None):
    """饼图：名称 + 数值 → 占比 PNG，嵌入 Word

    标签内嵌扇区（名称 + 百分比直接写在饼里），不靠图例猜；
    小于 4% 的碎块合并为"其他"（不列明细，保持饼图干净）。
    数据准确性：raw_values（与 values 对应的原始份额）非空时，图内百分比
    标注原始份额（与表格一致），不再显示归一化重算值（修复饼图/表格数字打架）。
    """
    import matplotlib.pyplot as plt
    if not labels or len(labels) != len(values) or not any(values):
        return
    total = sum(v for v in values if v > 0)
    if total <= 0:
        return
    keep = [(l, v) for l, v in zip(labels, values) if v > 0]
    keep.sort(key=lambda x: x[1], reverse=True)
    # 回归修复：raw 值按标签映射（原实现按原始输入顺序迭代，排序后扇区与标注错位；
    # "其他"合并项也没有对应 raw 值，会取到第一个小碎块的份额）
    raw_by_label = {}
    if raw_values and len(raw_values) == len(labels):
        raw_by_label = dict(zip(labels, raw_values))
    main = [(l, v) for l, v in keep if v / total >= 0.04]
    rest = [(l, v) for l, v in keep if v / total < 0.04]
    if rest:
        if raw_by_label:
            raw_by_label["其他"] = sum(raw_by_label.get(l, 0) for l, _ in rest)
        main.append(("其他", sum(v for _, v in rest)))
    if not main:
        main = keep[:6]  # 全都很小时兜底取前 6

    fig, ax = plt.subplots(figsize=(9.5, 5.6))
    labels_pie = [l for l, _ in main]
    values_pie = [v for _, v in main]
    colors = ["#2e5bff", "#e06a4c", "#12b886", "#f0a020", "#8e6fc0", "#5aa7d4", "#c0c0c0"]
    # 原始份额按扇区顺序（labels_pie 与扇区一一对应；缺失回退归一化百分比）
    _raw_iter = iter(labels_pie)

    def _autopct(p):
        try:
            label = next(_raw_iter)
        except StopIteration:
            return f"{p:.1f}%"
        if label in raw_by_label:
            return f"{raw_by_label[label]:.1f}%"
        return f"{p:.1f}%"

    wedges, texts, autotexts = ax.pie(
        values_pie, labels=None,
        autopct=_autopct,
        colors=colors[:len(main)], startangle=90,
        pctdistance=0.68,
        wedgeprops={"edgecolor": "white", "linewidth": 1.2})
    # 百分比内嵌（大字号白粗）；名称用引线拉到扇区外（leader line），字号加大
    import math
    for w, at, l in zip(wedges, autotexts, labels_pie):
        ang = math.radians((w.theta1 + w.theta2) / 2)
        # 百分比：内嵌扇区中心偏外，大字号
        at.set_position((0.72 * w.r * math.cos(ang), 0.72 * w.r * math.sin(ang)))
        at.set_fontsize(13)
        at.set_color("white")
        at.set_fontweight("bold")
        # 名称：引线拉到外圈，文字写在末端（加大字号）
        edge_x = 1.05 * w.r * math.cos(ang)
        edge_y = 1.05 * w.r * math.sin(ang)
        label_x = 1.35 * w.r * math.cos(ang)
        label_y = 1.35 * w.r * math.sin(ang)
        ax.plot([edge_x, label_x], [edge_y, label_y], color="#999", lw=0.9)
        ax.text(label_x, label_y, l, ha="center", va="center", fontsize=12.5,
                color="#333", fontweight="bold")
    # 标题放在图内顶部（加大画布，引线标签不与标题重叠）
    ax.set_title(title, fontsize=13, pad=18)
    ax.axis("equal")
    ax.set_xlim(-1.85, 1.85)
    ax.set_ylim(-1.5, 1.6)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    doc.add_picture(buf, width=Cm(13))


def build_csv(rows: list) -> io.BytesIO:
    """生成 CSV 原始数据：完整导出 UN Comtrade 返回的每条记录（所有字段）

    空数据返回带表头的空文件（标准 CSV），不再返回 "暂无数据" 文本（审查 #17）。
    """
    # 取所有记录并集字段（保持顺序），UN 返回啥导啥
    all_keys: list[str] = []
    seen = set()
    for r in rows:
        if isinstance(r, dict):
            for k in r.keys():
                if k not in seen:
                    seen.add(k)
                    all_keys.append(k)

    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(all_keys)
    for r in rows:
        if isinstance(r, dict):
            writer.writerow([r.get(k, "") for k in all_keys])
    return io.BytesIO(buf.getvalue().encode("utf-8-sig"))  # BOM 防 Excel 中文乱码


def build_agent_report(md: str, product: str, country: str) -> io.BytesIO:
    """AI Agent 的 markdown 报告 → 学术式 Word（复用字体/封面/收尾体系）

    markdown 按行解析：# 标题、## 子标题、- 列表、> 引用、普通段落。
    封面复用 build_market_report 的品牌块 + 信息行；正文/标题字体与正式报告一致。
    """
    from docx.shared import RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    _apply_doc_fonts(doc)
    style = doc.styles["Normal"]
    style.font.name = FONT_BODY
    style.font.size = Pt(12)

    NAVY = RGBColor(0x12, 0x3C, 0x5C)
    ACCENT = RGBColor(0xC4, 0x45, 0x2C)
    MUTED = RGBColor(0x7A, 0x71, 0x5F)

    # ===== 封面（品牌行 + 装饰线 + 标题 + 信息块）=====
    p0 = doc.add_paragraph()
    r0 = p0.add_run("TradePilot AI · EXPORT INTELLIGENCE")
    r0.font.name = "Arial"
    r0._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    r0.font.size = Pt(10)
    r0.font.color.rgb = ACCENT
    r0.bold = True

    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(6)
    hr.paragraph_format.space_after = Pt(18)
    pPr = hr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C4452C")
    pBdr.append(bottom)
    pPr.append(pBdr)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(f"{product} 市场分析报告（{country}）")
    tr.font.name = FONT_HEADING
    tr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
    tr.font.size = Pt(20)
    tr.bold = True
    tr.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("AI Agent 一句话全流程生成 · TradePilot AI")
    sr.font.size = Pt(10)
    sr.font.color.rgb = MUTED

    info = doc.add_paragraph()
    info.paragraph_format.space_before = Pt(24)
    ir = info.add_run(f"生成日期：{datetime.date.today().isoformat()}    数据来源：UN Comtrade · World Bank · Tavily 行业检索\n统计指标由程序精确计算 · AI 估算处明确标注")
    ir.font.size = Pt(9)
    ir.font.color.rgb = MUTED
    doc.add_page_break()

    # ===== 正文：markdown → 段落 =====
    def _h(text, level=1):
        return doc.add_heading(text, level=level)

    def _p(text=""):
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(24)
        r = para.add_run(text)
        return para

    def _li(text):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Pt(18)
        r = para.add_run("• " + text)
        return para

    for line in md.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith("### "):
            _h(line[4:], level=2)
        elif line.startswith("## "):
            _h(line[3:], level=1)
        elif line.startswith("# "):
            _h(line[2:], level=0)
        elif line.startswith("> "):
            para = doc.add_paragraph()
            r = para.add_run(line[2:])
            r.font.color.rgb = MUTED
            r.italic = True
        elif line.startswith("- ") or line.startswith("* "):
            _li(line[2:])
        elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. "):
            _li(line[3:])
        else:
            _p(line)

    _buf = io.BytesIO()
    doc.save(_buf)
    # 不在函数内 finalize：路由统一调 finalize_docx（与 build_market_report 一致），
    # 避免重复跑 COM 域更新
    return _buf

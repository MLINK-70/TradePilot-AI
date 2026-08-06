"""生成 TradePilot AI 专业报告模板 v2（docxtpl）

统一设计：
- 正文: 微软雅黑 11pt，深灰
- 标题1: 深蓝 16pt
- 标题2: 深蓝 13pt
- 图表坑位 + 表格坑位（docxtpl 渲染）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 全局默认样式：微软雅黑 11pt
style = doc.styles["Normal"]
style.font.name = "微软雅黑"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
pf = style.paragraph_format
pf.space_after = Pt(6)
pf.line_spacing = 1.4

# 标题1样式
h1 = doc.styles["Heading 1"]
h1.font.name = "微软雅黑"
h1.font.size = Pt(16)
h1.font.bold = True
h1.font.color.rgb = RGBColor(0x2E, 0x5B, 0xFF)
h1.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

# 标题2样式
h2 = doc.styles["Heading 2"]
h2.font.name = "微软雅黑"
h2.font.size = Pt(13)
h2.font.bold = True
h2.font.color.rgb = RGBColor(0x1A, 0x23, 0x33)
h2.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

# ===== 报告封面 =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TradePilot AI")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = RGBColor(0x2E, 0x5B, 0xFF)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("贸易数据分析报告")
run2.font.size = Pt(15)
run2.font.color.rgb = RGBColor(0x1A, 0x23, 0x33)

# 元信息（docxtpl 坑位）
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("{{ meta_line }}")
mr.font.size = Pt(10)
mr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()  # 空行

# 一、出口趋势（图坑）
doc.add_heading("一、出口趋势", level=1)
doc.add_paragraph("{{ trend_image }}")

# 二、数据总览（表坑）
doc.add_heading("二、数据总览", level=1)
doc.add_paragraph("{{ overview_table }}")

# 三、原始数据（表坑）
doc.add_heading("三、原始数据（UN Comtrade）", level=1)
doc.add_paragraph("{{ data_table }}")

# 四、AI 市场分析
doc.add_heading("四、AI 市场分析", level=1)
doc.add_paragraph("{{ analysis_text }}")

# 声明
doc.add_paragraph()
f = doc.add_paragraph()
fr = f.add_run("数据来源：UN Comtrade 公共 API ｜ 贸易数据为官方汇总，AI 分析由大模型生成，仅供参考")
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.save("templates/report_template_v2.docx")
print("模板 v2 已生成")

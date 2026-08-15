# -*- coding: utf-8 -*-
"""生成给老师演示的两份材料（零 API 消耗）：
1. TradePilot-AI-演示报告样张-蓝牙耳机德国.docx —— 真实缓存报告 + 学术式排版
2. TradePilot-AI-项目介绍-老师版.docx —— 介绍模板（概述/架构/数据可信度/演示引导）
"""
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

import database
from export import build_agent_report, finalize_docx

DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


# ── 1. 演示报告样张（用 8/12 真实缓存）──────────────────────────────
row = database.get_report_history("market", "蓝牙耳机", "德国")
if not row:
    print("无缓存报告，跳过样张")
else:
    md = row.get("report", "")
    buf = build_agent_report(md, "蓝牙耳机", "德国")
    buf, fmt = finalize_docx(buf, as_pdf=False)
    out1 = os.path.join(DESKTOP, "TradePilot-AI-演示报告样张-蓝牙耳机德国.docx")
    with open(out1, "wb") as f:
        f.write(buf.getvalue())
    print(f"1. 报告样张已生成: {out1}（{fmt}，{len(md)} 字 markdown）")

# ── 2. 项目介绍（老师版）────────────────────────────────────────────
doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def h1(text):
    p = doc.add_paragraph()
    set_run(p.add_run(text), size=15, bold=True, color=RGBColor(0x1F, 0x3B, 0x73))
    p.paragraph_format.space_before = Pt(14)


def h2(text):
    p = doc.add_paragraph()
    set_run(p.add_run(text), size=12, bold=True)
    p.paragraph_format.space_before = Pt(8)


def para(text, bold=False):
    p = doc.add_paragraph()
    set_run(p.add_run(text), size=10.5, bold=bold)
    p.paragraph_format.space_after = Pt(4)


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run(text), size=10.5)


t = doc.add_paragraph()
t.alignment = 1
set_run(t.add_run("TradePilot AI — 跨境贸易智能平台"), size=20, bold=True, color=RGBColor(0x1F, 0x3B, 0x73))
st = doc.add_paragraph()
st.alignment = 1
set_run(st.add_run("项目介绍（面向导师）"), size=12, color=RGBColor(0x66, 0x66, 0x66))

h1("一、项目定位")
para("面向消费电子出海的中小外贸企业，提供从市场调研到客户开发的全流程 AI 辅助。"
     "核心形态：一句话输入（如「蓝牙耳机去德国卖」），平台自动完成 市场分析 → 报告 → 客户线索 → 定制开发信。")

h1("二、解决的实际问题")
bullet("中小外贸企业进入新市场，传统做法需购买市场报告（数千元/份）、雇人检索客户、手工撰写开发信，成本高、周期长。")
bullet("本平台将上述流程自动化：真实数据支撑的分析 + AI 解读 + 客户检索 + 开发信生成，分钟级完成。")

h1("三、系统架构")
para("Python · FastAPI · SQLite（WAL 模式）· 原生 HTML/JS · ECharts · matplotlib")
bullet("多 AI 提供商适配层（DeepSeek 默认 / GPT / Claude / 自定义），统一调用与缓存。")
bullet("数据层：UN Comtrade（贸易）、World Bank（经济）、Tavily（行业动态）、WTO（宏观背景）。")
bullet("模块：市场分析、贸易数据、外贸业务、跨境电商、客户线索、AI Agent 编排、管理面板。")
bullet("部署：Dockerfile（含 LibreOffice 实现 Linux PDF 导出）+ docker-compose + GitHub Actions CI。")

h1("四、数据可信度（答辩重点）")
para("市场分析类 AI 应用的共性问题：AI 生成内容可能失真。本项目设计三道防线：", bold=True)
bullet("① 算术与解读分离：所有统计指标（CAGR、竞争力指数 TC、市场份额）由程序精确计算，AI 仅负责解读与建议，不参与任何算术。")
bullet("② 证据链注入：分析前先聚合真实数据（贸易额、GDP、竞争格局）注入提示词，AI 必须基于给定数据回答；数据不足处强制标注「估算」。")
bullet("③ 防幻觉硬约束：客户线索必须携带来源 URL 且 URL 需出现在搜索结果中，否则剔除；报告可溯源至原始数据源。")

h1("五、工程化水平")
bullet("54 项自动化测试（pytest，统计指标/安全校验/缓存/LLM 重试/导出，无网络依赖）。")
bullet("四路并行代码审查驱动：修复 60+ 问题（安全基线、缓存回归、并发竞态等）。")
bullet("性能：缓存 TTL 体系 + 并发证据链采集，欧盟 27 国聚合从串行 30s+ 优化为并发。")
bullet("安全：CORS 收紧、SSRF 防线、提示词注入防护、管理员权限层、匿名限流。")

h1("六、演示引导（约 3 分钟）")
bullet("① 首页/导航进入 AI Agent 页，输入「蓝牙耳机去德国卖」，展示六步流水线实时进度。")
bullet("② 展示生成的报告（可下载 Word/PDF 学术式排版）与客户线索（带来源 URL）。")
bullet("③ 打开管理面板，展示安全拦截记录（未授权访问被拒绝的实时日志）。")
bullet("④ 快速演示贸易数据页：HS 编码查询真实出口数据与趋势图。")

h1("七、数据来源声明")
para("演示数据来自 UN Comtrade 公共贸易数据库、World Bank 开放数据；行业动态来自 Tavily 网页检索；"
     "AI 输出基于上述证据链生成，估算处明确标注。")

out2 = os.path.join(DESKTOP, "TradePilot-AI-项目介绍-老师版.docx")
doc.save(out2)
print(f"2. 项目介绍已生成: {out2}")
